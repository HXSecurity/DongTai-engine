#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author: luzhongyang@huoxian.cn
# datetime: 2021/11/03 下午2:26
# project: DongTai-engine
from dongtai.models.project import IastProject
import time
from dongtai.models.vulnerablity import IastVulnerabilityModel
from collections import namedtuple
from django.db.models import Q
from dongtai.models.hook_type import HookType
from dongtai.models.agent import IastAgent
from django.utils.translation import gettext_lazy as _
from django.utils.translation import override
import re
import json
from dongtai.models.vul_level import IastVulLevel
from dongtai.models.message import IastMessage, IastMessageType
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from lingzhi_engine.settings import MEDIA_ROOT
from io import BytesIO
from django.utils.translation import gettext as _

import os


def get_model_field(model, exclude=[], include=[]):
    fields = [field.name for field in model._meta.fields]
    if include:
        return [
            field for field in list(set(fields) - set(exclude))
            if field in include
        ]
    return list(set(fields) - set(exclude))


def delete_old_files(path, save_seconds=10):
    for f in os.listdir(path):
        if os.stat(os.path.join(path, f)).st_mtime < time.time() - save_seconds:
            os.remove(os.path.join(path, f))


def get_vul_count_by_agent(agent_ids, vid, user):
    typeInfo = IastVulnerabilityModel.objects.filter(
        agent_id__in=agent_ids).values().order_by("level")
    if vid:
        typeInfo = typeInfo.filter(id=vid)
    type_summary = []
    levelCount = {}
    vulDetail = {}
    if typeInfo:
        typeArr = {}
        typeLevel = {}
        for one in typeInfo:
            hook_type = HookType.objects.filter(pk=one['hook_type_id']).first()
            one['type'] = hook_type.name if hook_type else ''
            typeArr[one['type']] = typeArr.get(one['type'], 0) + 1
            typeLevel[one['type']] = one['level_id']
            levelCount[one['level_id']] = levelCount.get(one['level_id'], 0) + 1
            language = IastAgent.objects.filter(
                pk=one['agent_id']).values_list('language', flat=True).first()
            one['language'] = language if language is not None else ''
            if one['type'] not in vulDetail.keys():
                vulDetail[one['type']] = []
            detailStr1 = _(
                "We found that there is {1} in the {0} page, attacker can modify the value of {2} to attack:").format(
                one['uri'], one['type'], one['taint_position'])

            try:
                one['req_params'] = str(one['req_params'])
            except Exception as e:
                one['req_params'] = ""
            detailStr2 = one['http_method'] + " " + one['uri'] + "?" + one['req_params'] + one['http_protocol']
            try:
                fileData = one['full_stack'][-1].get("stack", "")
                pattern = r'.*?\((.*?)\).*?'
                resMatch = re.match(pattern, fileData)
                uriArr = resMatch.group(1).split(":")
                fileName = uriArr[0]
                if len(uriArr) > 1:
                    rowStr = _("{} Line").format(str(uriArr[1]))
                else:
                    rowStr = ""
            except Exception as e:
                fileName = ""
                rowStr = ""
            classname = ""
            methodname = ""
            if one['full_stack']:
                try:
                    full_stack_arr = json.loads(one['full_stack'])
                    full_stack = full_stack_arr[-1]
                    classname = str(full_stack.get("classname", ""))
                    methodname = str(full_stack.get("methodname", ""))
                except Exception as e:
                    print("======")
            detailStr3 = _("In {} {} call {}. {} (), Incoming parameters {}").format(
                str(fileName), rowStr, classname, methodname,
                str(one['taint_value']))
            cur_tile = _("{} Appears in {} {}").format(one['type'], str(one['uri']), str(one['taint_position']))
            if one['param_name']:
                cur_tile = cur_tile + "\"" + str(one['param_name']) + "\""
            vulDetail[one['type']].append({
                "title": cur_tile,
                "type_name": one['type'],
                "level_id": one['level_id'],
                "first_time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(one['first_time'])),
                "latest_time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(one['latest_time'])),
                "language": one['language'],
                "url": one['url'],
                "detail_data": [detailStr1, detailStr2, detailStr3],
            })
        typeArrKeys = typeArr.keys()
        for item_type in typeArrKeys:
            type_summary.append(
                {
                    'type_name': item_type,
                    'type_count': typeArr[item_type],
                    'type_level': typeLevel[item_type]
                }
            )
    return {
        'type_summary': type_summary,
        'levelCount': levelCount,
        'vulDetail': vulDetail
    }


def get_translation_in(language, s):
    with override(language):
        return _(s)


class ExportPort():

    def export(self, report):
        # print(_("Department does not exist"))
        if report:
            self.generate_report(report)

    def get_agents_with_project_id(self, pid):
        """
        :param pid:
        :param auth_users:
        :return:
        """
        relations = IastAgent.objects.filter(bind_project_id=pid).values("id")
        agent_ids = [relation['id'] for relation in relations]
        return agent_ids

    def generate_report(self, report):
        type = report.type
        pid = 0
        if report.project:
            pid = report.project.id
        vid = report.vul_id
        user = report.user
        timestamp = time.time()

        project = IastProject.objects.filter(Q(id=pid)).first()

        vul = IastVulnerabilityModel.objects.filter(pk=vid).first()

        if project or vul:
            if not project:
                Project = namedtuple(
                    'Project',
                    get_model_field(IastProject,
                                    include=[
                                        'id', 'name', 'mode', 'latest_time',
                                        'vul_count', 'agent_count'
                                    ]))
                project = Project(id=0,
                                  name='NAN',
                                  mode='NAN',
                                  latest_time=time.time(),
                                  vul_count=1,
                                  agent_count=0)
            agent_ids = self.get_agents_with_project_id(project.id)

            count_result = get_vul_count_by_agent(agent_ids, vid, user)

            levelInfo = IastVulLevel.objects.all()
            from django.utils.translation import activate
            if report.language:
                activate(report.language)
            else:
                activate("zh")
            file_path = ""
            if type == 'docx':
                file_path = self.generate_word_report(user, project, vul, count_result, levelInfo, timestamp)
            elif type == 'pdf':
                file_path = self.generate_pdf_report(user, project, vul, count_result, levelInfo, timestamp)
            elif type == 'xlsx':
                file_path = self.generate_xlsx_report(user, project, vul, count_result, levelInfo, timestamp)
            if file_path != "":
                bin_file = open(file_path, "rb")
                file_data = bin_file.read()
                bin_file.close()
                report.file = file_data
                report.status = 1
                report.save()
                IastMessage.objects.create(
                    message=_("Report export success"),
                    relative_url="/dede",
                    create_time=time.time(),
                    message_type=IastMessageType.objects.filter(pk=1).first(),
                    to_user_id=report.user.id,
                )

    def generate_word_report(self, user, project, vul, count_result, levelInfo, timestamp):
        document = Document()
        document.styles.add_style('TitleOne', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleTwo', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleThree', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'
        document.styles.add_style('TitleFour', WD_STYLE_TYPE.PARAGRAPH).font.name = 'Arial'

        document.add_heading(u'%s' % project.name, 0)

        document.add_heading(u'%s' % project.mode, 2)

        timeArray = time.localtime(project.latest_time)
        otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)

        pTime = document.add_paragraph(u'%s' % otherStyleTime)
        pTime.paragraph_format.space_before = Pt(400)
        pTime.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        pReport = document.add_paragraph(_(u'Security Testing Report'))
        pReport.paragraph_format.line_spacing = Pt(20)
        pReport.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        footer = document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.add_run(u'北京安全共识科技有限公司')
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_page_break()

        oneTitle = document.add_paragraph()
        oneTitle.add_run(_(u'First, project information')).font.name = 'Arial'
        oneTitle.style = "TitleOne"

        table = document.add_table(rows=1, cols=2, style='Table Grid')

        hdr_cells = table.rows[0].cells

        new_cells = table.add_row().cells
        new_cells[0].text = _('Application name')
        new_cells[1].text = project.name
        new_cells = table.add_row().cells
        new_cells[0].text = _('Author')
        new_cells[1].text = user.username
        new_cells = table.add_row().cells
        new_cells[0].text = _('Application type')
        new_cells[1].text = project.mode
        new_cells = table.add_row().cells
        new_cells[0].text = _('Number of Vulnerability')
        new_cells[1].text = str(project.vul_count)
        new_cells = table.add_row().cells
        new_cells[0].text = _('Number of Agent')
        new_cells[1].text = str(project.agent_count)
        new_cells = table.add_row().cells
        new_cells[0].text = _('Latest time')
        new_cells[1].text = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)

        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        oneTitle = document.add_paragraph()
        oneTitle.add_run(_(u'Second, the result analysis'))
        oneTitle.style = "TitleOne"

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.1 Vulnerability Severity Levels Distribution'))
        twoTitle.style = "TitleTwo"
        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)
        document.add_paragraph(levelCountStr)

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.2 Distribution of Vulnerability'))
        twoTitle.style = "TitleTwo"

        table = document.add_table(rows=1, cols=3, style='Table Grid')

        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = _('Severity levels')
        hdr_cells[1].text = _('Vulnerability type name')
        hdr_cells[2].text = _('Number')
        if type_summary:
            for type_item in type_summary:
                new_cells = table.add_row().cells
                new_cells[0].text = levelIdArr[type_item['type_level']]
                new_cells[1].text = type_item['type_name']
                new_cells[2].text = str(type_item['type_count'])

        document.add_page_break()

        twoTitle = document.add_paragraph()
        twoTitle.add_run(_(u'2.3 Vulnerability details'))
        twoTitle.style = "TitleTwo"

        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():

                threeTitle = document.add_paragraph()
                threeTitle.add_run(u'%s(%s)' % ("2.3." + str(type_ind) + "  " + vul, len(vulDetail[vul])))
                threeTitle.style = "TitleThree"
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        p = document.add_paragraph()
                        p.add_run("2.3." + str(type_ind) + "." + str(ind) + "  " + one['title']).bold = True
                        p.style = "TitleFour"
                        ind = ind + 1
                        document.add_heading(_(u'Summary'), level=4)

                        table = document.add_table(rows=1, cols=2, style='Table Grid')
                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Severity level")
                        new_cells[1].text = levelIdArr[one['level_id']]

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("First scan time")
                        new_cells[1].text = one['first_time']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Last scan time")
                        new_cells[1].text = one['latest_time']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Development language")
                        new_cells[1].text = one['language']

                        new_cells = table.add_row().cells
                        new_cells[0].text = _("Vulnerability URL")
                        new_cells[1].text = one['url']
                        document.add_heading(_(u'Vulnerability description'), level=4)
                        if one['detail_data']:
                            for item in one['detail_data']:
                                document.add_paragraph(u'%s' % item)
                type_ind = type_ind + 1

        document.styles['TitleOne'].font.size = Pt(20)
        document.styles['TitleOne'].font.name = "Arial"
        document.styles['TitleTwo'].font.size = Pt(18)
        document.styles['TitleTwo'].font.name = "Arial"
        document.styles['TitleThree'].font.size = Pt(16)
        document.styles['TitleFour'].font.size = Pt(14)
        filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.docx"
        document.save(filename)

        return filename

    def generate_pdf_report(self, user, project, vul, count_result, levelInfo, timestamp):
        from django.template.loader import render_to_string
        import os

        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)

        vulTypeTableBodyRows = []

        new_cells = []
        if type_summary:
            for type_item in type_summary:
                vulTypeTableBodyRow = {
                    "type_level": levelIdArr[type_item['type_level']],
                    "type_name": type_item['type_name'],
                    "type_count": str(type_item['type_count'])
                }
                vulTypeTableBodyRows.append(vulTypeTableBodyRow)

        vulTypeDetailArray = []
        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():
                vulTypeDetail = {
                    "title": u'%s(%s)' % ("2.3." + str(type_ind) + "  " + vul, len(vulDetail[vul])),
                    "vuls": []
                }
                vulTypeDetailArray.append(vulTypeDetail)
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        oneVul = {
                            "title": "2.3." + str(type_ind) + "." + str(ind) + "  " + one['title'],
                            "summary": _(u'Summary'),

                            "severity_level": _("Severity level"),
                            "level_id": levelIdArr[one['level_id']],
                            "first_scan_time": _("First scan time"),
                            "first_time": one['first_time'],
                            "last_scan_time": _("First scan time"),
                            "latest_time": one['first_time'],
                            "development_language": _("Development language"),
                            "language": one['language'],
                            "vulnerability_url": _("Vulnerability URL"),
                            "url": one['url'],

                            "description": _(u'Vulnerability description'),
                            "detail": "",
                        }
                        vulTypeDetail.vuls.append(
                            oneVul
                        )
                        ind = ind + 1
                        if one['detail_data']:
                            for item in one['detail_data']:
                                oneVul.detail += u'%s' % item
                type_ind = type_ind + 1

        pdf_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.pdf"
        html_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.html"

        rendered = render_to_string(
            './pdf.html',
            {
                "user": user,
                "project": project,
                "vul": vul,
                "count_result": count_result,
                "level_info": levelInfo,
                "time_str": time.strftime('%Y-%m-%d %H:%M:%s', time.localtime(timestamp)),
                "levelCountStr": levelCountStr,
                "vulTypeTableBodyRows": vulTypeTableBodyRows,
                "i18n": {
                    "application_name": _("Application name"),
                    "author": _("Author"),
                    "number_of_vulnerability": _("Number of Vulnerability"),
                    "number_of_agent": _("Number of Agent"),
                    
                    "first_project_information": _("First, project information"),
                    "second_the_result_analysis": _("Second, the result analysis"),
                    "vulnerability_severity_levels_distribution": _("Vulnerability Severity Levels Distribution"),
                    "distribution_of_vulnerability": _("Distribution of Vulnerability"),
                    "severity_levels": _("Severity levels"),
                    "vulnerability_type_name": _("Vulnerability type name"),
                    "number": _("Number"),
                    "vulnerability_details": _("Vulnerability details"),
                }
            }
        )
        f = open(html_filename, 'w')
        f.write(rendered)
        f.close()
        os.system("cat {} | /usr/local/bin/wkhtmltopdf - {}".format(
            html_filename,
            pdf_filename
        ))
        delete_old_files(f"{MEDIA_ROOT}/reports/")
        return pdf_filename

    def generate_xlsx_report(self, user, project, vul, count_result, levelInfo, timestamp):
        levelNameArr = {}
        levelIdArr = {}
        if levelInfo:
            for level_item in levelInfo:
                levelNameArr[level_item.name_value] = level_item.id
                levelIdArr[level_item.id] = level_item.name_value

        type_summary = count_result['type_summary']

        levelCount = count_result['levelCount']

        vulDetail = count_result['vulDetail']

        levelCountArr = []
        if levelCount:
            for ind in levelCount.keys():
                levelCountArr.append(str(levelIdArr[ind]) + str(levelCount[ind]))
        levelCountStr = ",".join(levelCountArr)

        vulTypeTableBodyRows = []

        new_cells = []
        if type_summary:
            for type_item in type_summary:
                vulTypeTableBodyRow = {
                    "type_level": levelIdArr[type_item['type_level']],
                    "type_name": type_item['type_name'],
                    "type_count": str(type_item['type_count'])
                }
                vulTypeTableBodyRows.append(vulTypeTableBodyRow)

        vulTypeDetailArray = []
        if vulDetail:
            type_ind = 1
            for vul in vulDetail.keys():
                vulTypeDetail = {
                    "title": u'%s' % ("2.3." + str(type_ind) + "  " + vul),
                    "vuls": []
                }
                vulTypeDetailArray.append(vulTypeDetail)
                if vulDetail[vul]:
                    ind = 1
                    for one in vulDetail[vul]:
                        oneVul = {
                            "title": "2.3." + str(type_ind) + "." + str(ind) + "  " + one['title'],
                            "summary": _(u'Summary'),

                            "severity_level": _("Severity level"),
                            "level_id": levelIdArr[one['level_id']],
                            "first_scan_time": _("First scan time"),
                            "first_time": one['first_time'],
                            "last_scan_time": _("Last scan time"),
                            "latest_time": one['last_time'],
                            "development_language": _("Development language"),
                            "language": one['language'],
                            "vulnerability_url": _("Vulnerability URL"),
                            "url": one['url'],

                            "description": _(u'Vulnerability description'),
                            "detail": "",
                        }
                        vulTypeDetail.vuls.append(
                            oneVul
                        )
                        ind = ind + 1
                        if one['detail_data']:
                            for item in one['detail_data']:
                                oneVul.detail += u'%s' % item
                type_ind = type_ind + 1
        from openpyxl import Workbook
        wb = Workbook()
        sheet1 = wb.active
        xlsx_filename = f"{MEDIA_ROOT}/reports/vul-report-{user.id}-{timestamp}.xlsx"

        sheet1['A1'] = str(_("Vulnerability type name"))
        sheet1['B1'] = str(_("Severity levels"))
        sheet1['C1'] = str(_("First scan time"))
        sheet1['D1'] = str(_("Last scan time"))
        sheet1['E1'] = str(_("Development language"))
        sheet1['F1'] = str(_("Vulnerability URL"))
        sheet1['G1'] = str(_('Vulnerability description'))
        line = 0
        for vulTypeDetail in vulTypeDetailArray:
            line += 1
            sheet1.write(line, 0, vulTypeDetail.title)
            for oneVul in vulTypeDetail.vuls:
                sheet1.append(
                    [vulTypeDetail.title, oneVul.level_id, oneVul.first_time, oneVul.latest_time, oneVul.language,
                     oneVul.url, oneVul.detail])
        wb.save(xlsx_filename)
        delete_old_files(f"{MEDIA_ROOT}/reports/")
        return xlsx_filename
